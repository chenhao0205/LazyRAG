"""Workflow-local Markdown and DOCX exporter for academic manuscripts."""

from __future__ import annotations

import json
import re
import uuid
from pathlib import Path
from typing import Any


HEADING = re.compile(r'^(#{1,6})\s+(.+?)\s*$')
BULLET = re.compile(r'^\s*[-*+]\s+(.+)$')
NUMBERED = re.compile(r'^\s*\d+[.)、]\s+(.+)$')


def _clean_inline(text: str) -> str:
    value = re.sub(r'!\[([^]]*)\]\([^)]+\)', r'\1', str(text or ''))
    value = re.sub(r'\[([^]]+)\]\([^)]+\)', r'\1', value)
    return re.sub(r'(\*\*|__|~~|`)', '', value).strip()


def _artifact_payload(path: str) -> Any:
    source = Path(str(path or '')).expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(f'Workflow input does not exist: {path}')
    text = source.read_text(encoding='utf-8')
    try:
        value = json.loads(text)
    except json.JSONDecodeError:
        return text
    if isinstance(value, dict):
        if 'data' in value:
            return value['data']
        if 'text' in value:
            return value['text']
    return value


def _json_object(value: Any) -> dict[str, Any]:
    current = value
    for _ in range(5):
        if isinstance(current, dict):
            if 'data' not in current:
                return current
            current = current['data']
            continue
        try:
            current = json.loads(str(current or '').strip())
        except (TypeError, json.JSONDecodeError):
            return {}
    return current if isinstance(current, dict) else {}


def _remote_inputs() -> dict[str, Any]:
    from lazymind.chat.engine.subagent.context import require_context

    values = require_context().params.get('remote_inputs') or {}
    if not isinstance(values, dict):
        raise RuntimeError('Workflow input paths are unavailable in this Attempt.')
    return values


def _run_root() -> Path:
    from lazymind.chat.engine.subagent.context import require_context

    workspace = str(require_context().workspace_path or '').strip()
    if not workspace:
        raise RuntimeError('The active Workflow workspace is unavailable.')
    root = Path(workspace) / 'academic-paper-delivery' / uuid.uuid4().hex
    root.mkdir(parents=True, exist_ok=True)
    return root


def _latest_manuscript(values: dict[str, Any]) -> str:
    for slot in ('second_revised_document', 'revised_document', 'draft_document'):
        path = values.get(slot)
        if isinstance(path, str) and path.strip():
            return str(_artifact_payload(path))
    raise ValueError('No manuscript artifact is available for export.')


def _safe_name(value: str, suffix: str) -> str:
    name = re.sub(r'[\\/:*?"<>|\r\n\t]+', '_', str(value or '')).strip(' ._')
    name = name[:100] or 'academic_paper'
    return name if name.lower().endswith(suffix) else name + suffix


def _title(markdown: str) -> str:
    match = re.search(r'^#\s+(.+?)\s*$', str(markdown or ''), re.MULTILINE)
    return _clean_inline(match.group(1)) if match else 'Academic Paper'


def _configure_document(document: Any) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
    normal = document.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    for level in range(1, 6):
        style = document.styles[f'Heading {level}']
        style.font.name = 'Times New Roman'
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '黑体')
        style.font.size = Pt(max(12, 18 - level))
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_page_numbers(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement('w:fldChar')
        begin.set(qn('w:fldCharType'), 'begin')
        instruction = OxmlElement('w:instrText')
        instruction.set(qn('xml:space'), 'preserve')
        instruction.text = ' PAGE '
        end = OxmlElement('w:fldChar')
        end.set(qn('w:fldCharType'), 'end')
        run._r.extend([begin, instruction, end])


def _table_block(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or '|' not in lines[start]:
        return None
    separator = lines[start + 1].strip()
    if not re.match(r'^\|?\s*:?-{3,}', separator):
        return None
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and '|' in lines[index] and lines[index].strip():
        if index != start + 1:
            rows.append([_clean_inline(cell) for cell in lines[index].strip().strip('|').split('|')])
        index += 1
    return rows, index


def _markdown_to_docx(markdown: str, output: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    _configure_document(document)
    lines = str(markdown or '').splitlines()
    index = 0
    first_heading = True
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        table = _table_block(lines, index)
        if table:
            rows, index = table
            if rows:
                width = max(len(row) for row in rows)
                doc_table = document.add_table(rows=len(rows), cols=width)
                doc_table.style = 'Table Grid'
                for row_index, row in enumerate(rows):
                    for column_index, value in enumerate(row):
                        doc_table.cell(row_index, column_index).text = value
            continue
        heading = HEADING.match(stripped)
        if heading:
            level = min(5, len(heading.group(1)))
            text = _clean_inline(heading.group(2))
            paragraph = document.add_heading(text, level=level)
            if first_heading and level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(18)
                first_heading = False
            index += 1
            continue
        bullet = BULLET.match(raw)
        numbered = NUMBERED.match(raw)
        if bullet:
            document.add_paragraph(_clean_inline(bullet.group(1)), style='List Bullet')
        elif numbered:
            document.add_paragraph(_clean_inline(numbered.group(1)), style='List Number')
        elif stripped.startswith('>'):
            paragraph = document.add_paragraph(_clean_inline(stripped.lstrip('> ')))
            paragraph.paragraph_format.left_indent = Pt(18)
        else:
            document.add_paragraph(_clean_inline(stripped))
        index += 1
    _add_page_numbers(document)
    document.save(str(output))


def compose_academic_paper_from_inputs() -> dict[str, Any]:
    """Export the latest approved manuscript from immutable Workflow input paths."""
    values = _remote_inputs()
    parameters = _json_object(_artifact_payload(str(values.get('generation_parameters') or '')))
    if not parameters:
        raise ValueError('generation_parameters must contain a JSON object.')
    markdown = _latest_manuscript(values).strip()
    if not markdown:
        raise ValueError('The selected academic manuscript is empty.')
    root = _run_root()
    title = _title(markdown)
    markdown_path = root / _safe_name(title, '.md')
    markdown_path.write_text(markdown + '\n', encoding='utf-8')
    output_format = str(parameters.get('output_format') or '').strip().lower()
    if output_format == 'md':
        final_path = markdown_path
        renderer = 'workflow_local_markdown_export'
    elif output_format == 'docx':
        final_path = root / _safe_name(title, '.docx')
        _markdown_to_docx(markdown, final_path)
        renderer = 'workflow_local_python_docx'
    else:
        raise ValueError('output_format must be md or docx.')
    return {
        'final_markdown': str(markdown_path),
        'final_paper': str(final_path),
        'metadata': {
            'title': title,
            'output_format': output_format,
            'citation_style': parameters.get('citation_style'),
            'paper_language': parameters.get('paper_language'),
            'renderer': renderer,
            'source_skill': parameters.get('source_skill'),
            'size': final_path.stat().st_size,
        },
    }
